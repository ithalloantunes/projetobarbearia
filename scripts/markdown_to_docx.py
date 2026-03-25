from __future__ import annotations

import re
import sys
import unicodedata
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT_DIR = Path(__file__).resolve().parents[1]

FIGURE_IMAGE_MAP: dict[str, tuple[Path, float]] = {
    "Figura 1 - Modelo EAP do Projeto": (ROOT_DIR / "docs_assets" / "figures" / "fig01_eap.png", 15.2),
    "Figura 2 - Espinha de Peixe: Atraso do Projeto": (ROOT_DIR / "docs_assets" / "figures" / "fig02_atraso.png", 15.2),
    "Figura 3 - Espinha de Peixe: Indisponibilidade do Sistema": (ROOT_DIR / "docs_assets" / "figures" / "fig03_indisponibilidade.png", 15.2),
    "Figura 4 - Gerenciamento de Recursos Humanos": (ROOT_DIR / "docs_assets" / "figures" / "fig04_rh.png", 15.2),
    "Figura 5 - Comunicação do Projeto": (ROOT_DIR / "docs_assets" / "figures" / "fig05_comunicacao.png", 15.2),
    "Figura 6 - Tela de Login": (ROOT_DIR / "docs_assets" / "figures" / "fig06_login.png", 15.0),
    "Figura 7 - Tela de Início": (ROOT_DIR / "docs_assets" / "figures" / "fig07_inicio.png", 14.4),
    "Figura 8 - Tela de Seleção de Serviços": (ROOT_DIR / "docs_assets" / "figures" / "fig08_servicos.png", 14.8),
    "Figura 9 - Tela de Seleção do Barbeiro": (ROOT_DIR / "docs_assets" / "figures" / "fig09_barbeiro.png", 14.8),
    "Figura 10 - Tela de Seleção de Data e Horário": (ROOT_DIR / "docs_assets" / "figures" / "fig10_data_horario.png", 14.8),
    "Figura 11 - Tela de Confirmação de Agendamento": (ROOT_DIR / "docs_assets" / "figures" / "fig11_confirmacao.png", 14.8),
    "Figura 12 - Tela de Visualizar Agenda do Barbeiro": (ROOT_DIR / "docs_assets" / "figures" / "fig12_agenda_barbeiro.png", 13.8),
    "Figura 13 - Tela de Visualizar Bloqueios de Agenda": (ROOT_DIR / "docs_assets" / "figures" / "fig13_bloqueios.png", 13.8),
    "Figura 14 - Tela de Visualizar Serviços": (ROOT_DIR / "docs_assets" / "figures" / "fig14_servicos.png", 14.8),
    "Figura 15 - Tela de Visualizar Barbeiros": (ROOT_DIR / "docs_assets" / "figures" / "fig15_barbeiros.png", 14.8),
    "Figura 16 - Tela do Painel Administrativo": (ROOT_DIR / "docs_assets" / "figures" / "fig16_admin.png", 14.8),
    "Figura 17 - Tela de Visualizar Agendamentos": (ROOT_DIR / "docs_assets" / "figures" / "fig17_agendamentos.png", 14.8),
    "Figura 18 - Tela de Relatório de Atendimentos Diários": (ROOT_DIR / "docs_assets" / "figures" / "fig18_relatorio.png", 14.8),
    "Figura 19 - Código do front-end com Next.js": (ROOT_DIR / "docs_assets" / "figures" / "fig19_frontend.png", 15.0),
    "Figura 20 - Código do back-end com Next.js API e Prisma": (ROOT_DIR / "docs_assets" / "figures" / "fig20_backend.png", 15.0),
}


HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
ORDERED_LIST_RE = re.compile(r"^\d+\.\s+(.*)$")
META_RE = re.compile(r"^\*\*(.+?):\*\*\s*(.*)$")


@dataclass
class Entry:
    text: str
    bookmark: str
    level: int = 1


@dataclass
class SectionBlock:
    title: str
    lines: list[str]


def slugify(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    ascii_text = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    ascii_text = re.sub(r"[^A-Za-z0-9]+", "_", ascii_text).strip("_")
    return ascii_text.lower() or "item"


def clean_markdown_text(text: str) -> str:
    text = text.replace("**", "")
    return text.strip()


def ensure_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = name
    run.font.size = Pt(size)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_field_run(paragraph, instruction: str, default_text: str = "1", font_name: str = "Times New Roman", font_size: int = 12):
    begin = paragraph.add_run()
    ensure_font(begin, font_name, font_size)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin._r.append(fld_begin)

    instr = paragraph.add_run()
    ensure_font(instr, font_name, font_size)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr._r.append(instr_text)

    separate = paragraph.add_run()
    ensure_font(separate, font_name, font_size)
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separate._r.append(fld_separate)

    result = paragraph.add_run(default_text)
    ensure_font(result, font_name, font_size)

    end = paragraph.add_run()
    ensure_font(end, font_name, font_size)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end._r.append(fld_end)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def parse_sections(lines: list[str]) -> list[SectionBlock]:
    sections: list[SectionBlock] = []
    current_title: str | None = None
    current_lines: list[str] = []

    for raw_line in lines:
        stripped = raw_line.strip().lstrip("\ufeff")
        match = HEADING_RE.match(stripped)
        if match and len(match.group(1)) == 1:
            if current_title is not None:
                sections.append(SectionBlock(title=current_title, lines=current_lines))
            current_title = match.group(2).strip()
            current_lines = []
            continue

        if current_title is not None:
            current_lines.append(raw_line)

    if current_title is not None:
        sections.append(SectionBlock(title=current_title, lines=current_lines))

    return sections


def parse_cover_data(section: SectionBlock) -> tuple[str, dict[str, str]]:
    subtitle = ""
    metadata: dict[str, str] = {}

    for line in section.lines:
        stripped = line.strip()
        heading_match = HEADING_RE.match(stripped)
        meta_match = META_RE.match(stripped)

        if heading_match and len(heading_match.group(1)) == 2:
            subtitle = heading_match.group(2).strip()
        elif meta_match:
            metadata[meta_match.group(1).strip()] = meta_match.group(2).strip()

    return subtitle, metadata


def collect_entries(section: SectionBlock) -> list[str]:
    entries: list[str] = []
    for line in section.lines:
        stripped = line.strip()
        if not stripped or stripped == "---":
            continue
        if HEADING_RE.match(stripped):
            continue
        if META_RE.match(stripped):
            continue
        entries.append(clean_markdown_text(stripped))
    return entries


def extract_body_lines(sections: list[SectionBlock]) -> list[str]:
    body_lines: list[str] = []
    for section in sections[5:]:
        body_lines.append(f"# {section.title}")
        body_lines.extend(section.lines)
    return body_lines


def collect_body_indexes(body_lines: list[str]) -> tuple[list[Entry], list[Entry], list[Entry]]:
    toc_entries: list[Entry] = []
    figure_entries: list[Entry] = []
    table_entries: list[Entry] = []

    toc_count = 0
    fig_count = 0
    table_count = 0

    for raw_line in body_lines:
        stripped = raw_line.strip()
        heading_match = HEADING_RE.match(stripped)
        if stripped.startswith("Tabela "):
            table_count += 1
            table_entries.append(Entry(text=stripped, bookmark=f"tab_{table_count}_{slugify(stripped)}"))
            continue
        if not heading_match:
            continue

        level = len(heading_match.group(1))
        text = heading_match.group(2).strip()

        if text.startswith("Figura "):
            fig_count += 1
            figure_entries.append(Entry(text=text, bookmark=f"fig_{fig_count}_{slugify(text)}"))
            continue

        if text.startswith("Tabela "):
            table_count += 1
            table_entries.append(Entry(text=text, bookmark=f"tab_{table_count}_{slugify(text)}"))
            continue

        toc_count += 1
        toc_entries.append(Entry(text=text, bookmark=f"toc_{toc_count}_{slugify(text)}", level=level))

    return toc_entries, figure_entries, table_entries


def configure_page(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)


def configure_styles(document: Document) -> None:
    configure_page(document.sections[0])

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(24)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "List Paragraph" in document.styles:
        style = document.styles["List Paragraph"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

    if "Caption" not in document.styles:
        document.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)

    caption = document.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.bold = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = 1.0


def enable_field_update_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_inline_runs(paragraph, text: str, font_size: int = 12) -> None:
    index = 0
    while index < len(text):
        if text.startswith("**", index):
            end = text.find("**", index + 2)
            if end != -1:
                run = paragraph.add_run(text[index + 2 : end])
                ensure_font(run, "Times New Roman", font_size, bold=True)
                index = end + 2
                continue

        if text.startswith("`", index):
            end = text.find("`", index + 1)
            if end != -1:
                run = paragraph.add_run(text[index + 1 : end])
                ensure_font(run, "Courier New", 10)
                index = end + 1
                continue

        next_positions = [pos for pos in (text.find("**", index), text.find("`", index)) if pos != -1]
        next_stop = min(next_positions) if next_positions else len(text)
        run = paragraph.add_run(text[index:next_stop])
        ensure_font(run, "Times New Roman", font_size)
        index = next_stop


def add_blank_lines(document: Document, amount: int) -> None:
    for _ in range(amount):
        document.add_paragraph("")


def add_centered_text(document: Document, text: str, bold: bool = True, size: int = 12) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    ensure_font(run, "Times New Roman", size, bold=bold)


def add_right_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(9)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    add_inline_runs(paragraph, text, font_size=12)


def build_cover(document: Document, title: str, subtitle: str, metadata: dict[str, str]) -> None:
    institution = metadata.get("Instituição", "Faculdade de Tecnologia (FATEC)").upper()
    project = metadata.get("Projeto", "BarberSaaS").upper()
    location = metadata.get("Local", "São Paulo - SP").upper()
    year = metadata.get("Ano", "2026")
    members = [item.strip() for item in metadata.get("Integrantes", "").split(";") if item.strip()]

    add_blank_lines(document, 2)
    add_centered_text(document, institution, bold=True, size=12)
    add_blank_lines(document, 4)

    if members:
        for member in members:
            add_centered_text(document, member.upper(), bold=True, size=12)
    else:
        add_centered_text(document, "EQUIPE DO PROJETO BARBERSAAS", bold=True, size=12)

    add_blank_lines(document, 8)
    add_centered_text(document, title.upper(), bold=True, size=12)
    add_blank_lines(document, 1)

    subtitle_parts = subtitle.upper().split(" EM ")
    for index, chunk in enumerate(subtitle_parts):
        if chunk:
            line = chunk if index == 0 else f"EM {chunk}"
            add_centered_text(document, line, bold=True, size=12)

    add_blank_lines(document, 2)
    add_centered_text(document, project, bold=True, size=13)
    add_blank_lines(document, 10)
    add_centered_text(document, location, bold=True, size=12)
    add_centered_text(document, year, bold=True, size=12)


def build_title_page(document: Document, subtitle: str, metadata: dict[str, str], section: SectionBlock) -> None:
    institution = metadata.get("Instituição", "Faculdade de Tecnologia (FATEC)").upper()
    project = metadata.get("Projeto", "BarberSaaS")
    location = metadata.get("Local", "São Paulo - SP").upper()
    year = metadata.get("Ano", "2026")
    members = [item.strip() for item in metadata.get("Integrantes", "").split(";") if item.strip()]
    nature = metadata.get(
        "Natureza do trabalho",
        "Documentação acadêmico-técnica de projeto de desenvolvimento de software",
    )

    paragraphs = [clean_markdown_text(line.strip()) for line in section.lines if line.strip() and not HEADING_RE.match(line.strip()) and not META_RE.match(line.strip())]
    description = " ".join(paragraphs[:2]).strip()
    if not description:
        description = (
            "Documentação técnica elaborada como registro formal do projeto BarberSaaS, "
            "voltado ao desenvolvimento de um sistema web para gestão de agendamentos em barbearias."
        )

    add_blank_lines(document, 1)
    add_centered_text(document, institution, bold=True, size=12)
    add_blank_lines(document, 2)
    if members:
        for member in members:
            add_centered_text(document, member.upper(), bold=True, size=12)
    else:
        add_centered_text(document, f"EQUIPE DO PROJETO {project.upper()}", bold=True, size=12)
    add_blank_lines(document, 4)
    add_centered_text(document, subtitle.upper(), bold=True, size=12)
    add_blank_lines(document, 6)
    add_right_block(
        document,
        (
            f"Trabalho apresentado à {metadata.get('Instituição', 'Faculdade de Tecnologia (FATEC)')}, "
            f"como {nature.lower()}, tomando como objeto de estudo e desenvolvimento o projeto {project}. "
            f"{description}\n\n"
            f"Projeto: {project}\n"
            f"Product Owner: {metadata.get('Product Owner', 'Dono da barbearia')}\n"
            f"Scrum Master: {metadata.get('Scrum Master', '')}\n"
            f"Equipe de Desenvolvimento: {metadata.get('Equipe de Desenvolvimento', '')}\n"
            f"Versão do documento: {metadata.get('Versão do documento', '1.0')}\n"
            f"Data de consolidação: {metadata.get('Data de consolidação', '16 de março de 2026')}"
        ),
    )
    add_blank_lines(document, 10)
    add_centered_text(document, location, bold=False, size=12)
    add_centered_text(document, year, bold=False, size=12)


def add_leader_entry(document: Document, text: str, bookmark: str | None = None, level: int = 1, bold: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(max(level - 1, 0) * 0.6)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    run = paragraph.add_run(text)
    ensure_font(run, "Times New Roman", 12, bold=bold)
    tab_run = paragraph.add_run("\t")
    ensure_font(tab_run, "Times New Roman", 12, bold=bold)

    if bookmark:
        add_field_run(paragraph, f"PAGEREF {bookmark} \\h", default_text=" ")
    else:
        spacer = paragraph.add_run(" ")
        ensure_font(spacer, "Times New Roman", 12, bold=bold)


def build_index_page(document: Document, title: str, entries: list[Entry]) -> None:
    add_blank_lines(document, 2)
    add_centered_text(document, title, bold=True, size=12)
    add_blank_lines(document, 2)
    for entry in entries:
        add_leader_entry(document, entry.text, bookmark=entry.bookmark, level=entry.level, bold=True)


def build_toc_page(document: Document, title: str) -> None:
    add_blank_lines(document, 2)
    add_centered_text(document, title, bold=True, size=12)
    add_blank_lines(document, 2)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_field_run(
        paragraph,
        'TOC \\o "1-4" \\h \\z \\u',
        default_text="Atualize o campo do sumario no Word, se necessario.",
    )


def add_page_number_header(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))

    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_field_run(paragraph, "PAGE", default_text=str(start))


def parse_table_row(line: str) -> list[str]:
    clean = line.strip()
    if clean.startswith("|"):
        clean = clean[1:]
    if clean.endswith("|"):
        clean = clean[:-1]
    return [cell.strip() for cell in clean.split("|")]


def is_table_delimiter(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and set(stripped) <= {"|", "-", ":", " "}


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        header_cells[idx].text = value
        header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in header_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if paragraph.runs:
                ensure_font(paragraph.runs[0], "Times New Roman", 10, bold=True)
        set_cell_shading(header_cells[idx], "D9E2F3")

    for row_values in rows[1:]:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].text = value
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in row_cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if paragraph.runs:
                    ensure_font(paragraph.runs[0], "Times New Roman", 10)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        ensure_font(run, "Courier New", 9)
        if idx < len(lines) - 1:
            run.add_break()


def flush_paragraph(document: Document, buffer: list[str]) -> None:
    if not buffer:
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline_runs(paragraph, " ".join(buffer).strip(), font_size=12)
    buffer.clear()


def add_reference_entry(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    add_inline_runs(paragraph, text.strip(), font_size=12)


def add_body_heading(document: Document, level: int, text: str, bookmark: str, bookmark_id: int) -> None:
    if text.startswith("Figura ") or text.startswith("Tabela "):
        paragraph = document.add_paragraph(style="Caption")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph = document.add_paragraph(style=f"Heading {min(level, 4)}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph.paragraph_format.first_line_indent = Cm(0)
    add_inline_runs(paragraph, text, font_size=12)
    add_bookmark(paragraph, bookmark, bookmark_id)


def add_table_caption(document: Document, text: str, bookmark: str, bookmark_id: int) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    add_inline_runs(paragraph, text, font_size=12)
    add_bookmark(paragraph, bookmark, bookmark_id)


def add_figure_image(document: Document, figure_text: str) -> None:
    image_info = FIGURE_IMAGE_MAP.get(figure_text)
    if not image_info:
        return

    image_path, width_cm = image_info
    if not image_path.exists():
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def render_body(document: Document, body_lines: list[str], toc_entries: list[Entry], figure_entries: list[Entry], table_entries: list[Entry]) -> None:
    paragraph_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code_block = False

    toc_index = 0
    fig_index = 0
    table_index = 0
    bookmark_id = 1
    line_index = 0
    first_primary_heading = True
    in_references_section = False

    while line_index < len(body_lines):
        line = body_lines[line_index]
        stripped = line.strip()

        if in_code_block:
            if stripped.startswith("```"):
                add_code_block(document, code_buffer)
                code_buffer.clear()
                in_code_block = False
            else:
                code_buffer.append(line)
            line_index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph(document, paragraph_buffer)
            in_code_block = True
            line_index += 1
            continue

        if not stripped or stripped == "---":
            flush_paragraph(document, paragraph_buffer)
            line_index += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph(document, paragraph_buffer)
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            if text.startswith("Figura "):
                entry = figure_entries[fig_index]
                fig_index += 1
            elif text.startswith("Tabela "):
                entry = table_entries[table_index]
                table_index += 1
            else:
                entry = toc_entries[toc_index]
                toc_index += 1
                in_references_section = text.upper() == "REFERÊNCIAS"

            if level == 1:
                if first_primary_heading:
                    first_primary_heading = False
                else:
                    document.add_page_break()

            add_body_heading(document, level, text, entry.bookmark, bookmark_id)
            bookmark_id += 1
            if text.startswith("Figura "):
                add_figure_image(document, text)
            line_index += 1
            continue

        if in_references_section:
            flush_paragraph(document, paragraph_buffer)
            add_reference_entry(document, stripped)
            line_index += 1
            continue

        if stripped.startswith("Tabela "):
            flush_paragraph(document, paragraph_buffer)
            entry = table_entries[table_index]
            table_index += 1
            add_table_caption(document, stripped, entry.bookmark, bookmark_id)
            bookmark_id += 1
            line_index += 1
            continue

        if stripped.startswith("|") and line_index + 1 < len(body_lines) and is_table_delimiter(body_lines[line_index + 1]):
            flush_paragraph(document, paragraph_buffer)
            table_lines = [stripped]
            line_index += 2
            while line_index < len(body_lines) and body_lines[line_index].strip().startswith("|"):
                table_lines.append(body_lines[line_index].strip())
                line_index += 1
            add_table(document, [parse_table_row(row) for row in table_lines])
            continue

        if stripped.startswith("- "):
            flush_paragraph(document, paragraph_buffer)
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline_runs(paragraph, stripped[2:].strip(), font_size=12)
            line_index += 1
            continue

        ordered_match = ORDERED_LIST_RE.match(stripped)
        if ordered_match:
            flush_paragraph(document, paragraph_buffer)
            paragraph = document.add_paragraph(style="List Number")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline_runs(paragraph, ordered_match.group(1).strip(), font_size=12)
            line_index += 1
            continue

        paragraph_buffer.append(stripped)
        line_index += 1

    flush_paragraph(document, paragraph_buffer)


def convert(markdown_path: Path, docx_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    sections = parse_sections(lines)

    if len(sections) < 6:
        raise ValueError("Estrutura inesperada do markdown. Esperadas páginas preliminares e corpo do texto.")

    cover_section = sections[0]
    title_page_section = sections[1]
    figures_section = sections[2]
    tables_section = sections[3]
    summary_section = sections[4]

    body_lines = extract_body_lines(sections)
    toc_entries, figure_entries, table_entries = collect_body_indexes(body_lines)

    cover_subtitle, metadata = parse_cover_data(cover_section)

    document = Document()
    configure_styles(document)
    enable_field_update_on_open(document)

    build_cover(document, cover_section.title, cover_subtitle, metadata)
    document.add_page_break()

    build_title_page(document, cover_subtitle, metadata, title_page_section)
    document.add_page_break()

    if figure_entries:
        build_index_page(document, figures_section.title, figure_entries)
    else:
        build_index_page(document, figures_section.title, [Entry(text=item, bookmark=None or "") for item in collect_entries(figures_section)])
    document.add_page_break()

    if table_entries:
        build_index_page(document, tables_section.title, table_entries)
    else:
        build_index_page(document, tables_section.title, [Entry(text=item, bookmark=None or "") for item in collect_entries(tables_section)])
    document.add_page_break()

    build_toc_page(document, summary_section.title)

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_section)
    add_page_number_header(body_section, start=5)

    render_body(document, body_lines, toc_entries, figure_entries, table_entries)
    document.save(docx_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Uso: python scripts/markdown_to_docx.py <entrada.md> <saida.docx>")
        return 1

    markdown_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])

    if not markdown_path.exists():
        print(f"Arquivo não encontrado: {markdown_path}")
        return 1

    convert(markdown_path, docx_path)
    print(f"Arquivo gerado: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
