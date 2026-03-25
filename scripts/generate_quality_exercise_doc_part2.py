from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import generate_quality_exercise_doc as base


ROOT = Path(__file__).resolve().parents[1]
base.OUT_DIR = ROOT / "docs_assets" / "exercicio_parte_2"
base.DOCX_PATH = ROOT / "EXERCICIO PARTE 2.docx"


def build_docx_part2(image_paths: dict[str, Path]) -> None:
    document = Document()
    base.configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Exercicio Parte 2")
    base.set_run_font(run, 18, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("BarberSaaS - Qualidade aplicada ao problema de comunicacao automatizada")
    base.set_run_font(run, 12, color=base.PRIMARY_DARK)

    base.add_body_paragraph(
        document,
        "Este segundo exercicio explora outro problema real do BarberSaaS: a fragilidade da comunicacao automatizada com o cliente. O recorte considera as notificacoes de agendamento, a redefinicao de senha por email e a dependencia de credenciais externas para que a experiencia do usuario seja completa.",
    )

    base.add_heading(document, "1. PDCAs aplicados ao problema de comunicacao", 1)

    base.add_heading(document, "1.1 PDCA - Confirmacao de agendamento ao cliente", 2)
    base.add_bullets(
        document,
        [
            "Plan: definir meta de notificar o cliente com dados corretos logo apos a reserva.",
            "Do: validar disparo de email e WhatsApp na confirmacao e no reagendamento.",
            "Check: conferir logs, respostas externas e percepcao do cliente sobre o recebimento.",
            "Act: ajustar fallbacks, mensagens e validacoes para reduzir falhas de comunicacao.",
        ],
    )
    base.add_image(document, image_paths["pdca_1"], "Figura 1 - PDCA para confirmacao de agendamento")

    base.add_heading(document, "1.2 PDCA - Recuperacao de senha por email", 2)
    base.add_bullets(
        document,
        [
            "Plan: garantir que o fluxo de redefinicao seja seguro, claro e confiavel.",
            "Do: validar geracao de token, expiracao e envio do link de redefinicao.",
            "Check: observar erros de email, expiracao e falhas percebidas pelo usuario.",
            "Act: reforcar mensagens, validacoes e rastreabilidade do fluxo de recuperacao.",
        ],
    )
    base.add_image(document, image_paths["pdca_2"], "Figura 2 - PDCA para recuperacao de senha")

    base.add_heading(document, "1.3 PDCA - Rastreabilidade e fallback das notificacoes", 2)
    base.add_bullets(
        document,
        [
            "Plan: reduzir a dependencia cega de integracoes externas e melhorar a observabilidade.",
            "Do: usar modos de teste, logs locais e checklist de credenciais.",
            "Check: revisar se as falhas sao detectadas rapidamente e se o fluxo principal continua funcionando.",
            "Act: padronizar monitoramento, orientacoes de ambiente e resposta a incidentes.",
        ],
    )
    base.add_image(document, image_paths["pdca_3"], "Figura 3 - PDCA para rastreabilidade das notificacoes")

    base.add_heading(document, "2. Diagramas de causa e efeito", 1)

    base.add_heading(document, "2.1 Cliente nao recebe a confirmacao do agendamento", 2)
    base.add_body_paragraph(
        document,
        "Neste diagrama, o efeito central e a quebra de confianca do cliente apos a reserva. Mesmo quando o agendamento foi salvo, a falta de retorno por email ou WhatsApp reduz a percepcao de seguranca e profissionalismo no sistema.",
    )
    base.add_image(document, image_paths["fishbone_1"], "Figura 4 - Causa e efeito para falha na confirmacao do agendamento")

    base.add_heading(document, "2.2 Cliente nao consegue redefinir a senha", 2)
    base.add_body_paragraph(
        document,
        "Aqui o foco esta na recuperacao de acesso. Se o link nao chega, expira cedo demais ou aponta para ambiente incorreto, o usuario pode ficar bloqueado do sistema e interpretar o produto como instavel ou inseguro.",
    )
    base.add_image(document, image_paths["fishbone_2"], "Figura 5 - Causa e efeito para falha na redefinicao de senha")

    base.add_heading(document, "2.3 Comunicacao automatizada sem rastreabilidade", 2)
    base.add_body_paragraph(
        document,
        "Este diagrama aborda o problema estrutural por tras das notificacoes: sem observabilidade, credenciais bem geridas e rotina de verificacao, a equipe demora para descobrir falhas e perde capacidade de agir antes que o cliente seja impactado.",
    )
    base.add_image(document, image_paths["fishbone_3"], "Figura 6 - Causa e efeito para falta de rastreabilidade nas notificacoes")

    base.add_heading(document, "3. Importancia das ferramentas", 1)
    base.add_body_paragraph(
        document,
        "O diagrama de causa e efeito ajuda a separar o sintoma do problema real. No caso das notificacoes, ele mostra que a falha nao nasce apenas da API externa, mas tambem de configuracao, processo, seguranca, dados e monitoramento. Isso orienta a equipe a corrigir a raiz do problema, e nao apenas o efeito visivel.",
    )
    base.add_body_paragraph(
        document,
        "O PDCA e importante porque transforma esse aprendizado em melhoria continua. Ele permite planejar como a comunicacao deve funcionar, executar os fluxos, checar resultados com evidencias e agir rapidamente para ajustar credenciais, mensagens, observabilidade e contingencia sem comprometer a experiencia principal do sistema.",
    )

    document.save(base.DOCX_PATH)


def main() -> None:
    base.ensure_dirs()

    image_paths = {
        "pdca_1": base.create_pdca_image(
            "pdca_confirmacao_agendamento.png",
            "PDCA 1 - Confirmacao de agendamento ao cliente",
            "Foco",
            "Entrega confiavel da mensagem apos a reserva",
            [
                "Definir meta de confirmar reserva por email ou WhatsApp.",
                "Garantir dados corretos de cliente, servico, horario e profissional.",
            ],
            [
                "Executar envio de notificacao na criacao e no reagendamento.",
                "Validar comportamento com e sem credenciais completas.",
            ],
            [
                "Conferir logs, respostas externas e retorno percebido pelo cliente.",
                "Medir falhas de entrega e inconsistencias de mensagem.",
            ],
            [
                "Ajustar textos, validacoes e caminhos de fallback.",
                "Documentar regras para manter a confirmacao confiavel.",
            ],
        ),
        "pdca_2": base.create_pdca_image(
            "pdca_recuperacao_senha.png",
            "PDCA 2 - Recuperacao de senha por email",
            "Foco",
            "Acesso seguro e recuperavel para o usuario",
            [
                "Definir prazo de expiracao, base URL e mensagem clara.",
                "Evitar exposicao de informacoes sensiveis no fluxo.",
            ],
            [
                "Gerar token, salvar hash e enviar link de redefinicao.",
                "Validar resposta generica e tratamento de usuarios inativos.",
            ],
            [
                "Observar falhas de email, token expirado e links incorretos.",
                "Conferir se o usuario entende o proximo passo.",
            ],
            [
                "Reforcar rastreabilidade, clareza de mensagem e seguranca do fluxo.",
                "Padronizar testes para evitar quebra futura.",
            ],
        ),
        "pdca_3": base.create_pdca_image(
            "pdca_rastreabilidade_notificacoes.png",
            "PDCA 3 - Rastreabilidade e fallback das notificacoes",
            "Foco",
            "Observabilidade da comunicacao automatizada",
            [
                "Definir indicadores de entrega, erro e modo degradado.",
                "Mapear dependencias externas e variaveis obrigatorias.",
            ],
            [
                "Usar logs, modo de teste e checklist de configuracao.",
                "Registrar comportamento quando email ou WhatsApp falham.",
            ],
            [
                "Revisar se a equipe descobre rapido os incidentes.",
                "Confirmar se o agendamento continua funcionando sem travar.",
            ],
            [
                "Padronizar monitoramento e resposta a falhas.",
                "Melhorar a preparacao do ambiente para cada entrega.",
            ],
        ),
        "fishbone_1": base.create_fishbone_image(
            "causa_efeito_confirmacao_cliente.png",
            "Diagrama 1 - Cliente nao recebe confirmacao do agendamento",
            "Sem confirmacao ao cliente",
            [
                ("Dados", ["email ou telefone incorretos", "dados incompletos do agendamento"]),
                ("Integracao", ["falha no envio externo", "payload invalido para mensagem"]),
                ("Configuracao", ["SMTP ausente", "token de WhatsApp invalido"]),
            ],
            [
                ("Processo", ["sem checklist de ambiente", "teste tardio das notificacoes"]),
                ("Monitoramento", ["sem alerta de falha", "logs pouco consultados"]),
                ("Experiencia", ["mensagem pouco clara", "cliente nao sabe se agendou"]),
            ],
        ),
        "fishbone_2": base.create_fishbone_image(
            "causa_efeito_reset_senha.png",
            "Diagrama 2 - Cliente nao consegue redefinir a senha",
            "Falha no reset de senha",
            [
                ("Token", ["expiracao curta ou vencida", "hash/token gerado com problema"]),
                ("Email", ["SMTP falhando", "caixa do cliente nao recebe a mensagem"]),
                ("Link", ["APP_BASE_URL incorreto", "url quebrada ou invalida"]),
            ],
            [
                ("Usuario", ["conta inativa", "email digitado errado"]),
                ("Seguranca", ["mensagem pouco orientativa", "fluxo sem rastreabilidade"]),
                ("Processo", ["sem teste periodico do reset", "sem evidencia de entrega"]),
            ],
        ),
        "fishbone_3": base.create_fishbone_image(
            "causa_efeito_rastreabilidade.png",
            "Diagrama 3 - Comunicacao automatizada sem rastreabilidade",
            "Baixa rastreabilidade",
            [
                ("Observabilidade", ["logs dispersos", "sem indicador de entrega"]),
                ("Ambiente", ["credenciais mudam sem controle", "diferenca entre local e producao"]),
                ("Dependencias", ["servicos externos instaveis", "falha sem retorno claro"]),
            ],
            [
                ("Operacao", ["reacao lenta a incidentes", "sem rotina de verificacao"]),
                ("Qualidade", ["validacao manual insuficiente", "teste so no fim da entrega"]),
                ("Gestao", ["prioridade baixa para monitoramento", "documentacao de suporte incompleta"]),
            ],
        ),
    }

    build_docx_part2(image_paths)
    print(f"Imagens geradas em: {base.OUT_DIR}")
    print(f"Arquivo Word gerado em: {base.DOCX_PATH}")


if __name__ == "__main__":
    main()
