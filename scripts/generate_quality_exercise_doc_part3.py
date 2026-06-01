from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import generate_quality_exercise_doc as base


ROOT = Path(__file__).resolve().parents[1]
base.OUT_DIR = ROOT / "docs_assets" / "exercicio_parte_3"
base.DOCX_PATH = ROOT / "EXERCICIO PARTE 3.docx"


def build_docx_part3(image_paths: dict[str, Path]) -> None:
    document = Document()
    base.configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Exercicio Parte 3")
    base.set_run_font(run, 18, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("BarberSaaS - Qualidade aplicada a autenticacao e controle de permissoes")
    base.set_run_font(run, 12, color=base.PRIMARY_DARK)

    base.add_body_paragraph(
        document,
        "Esta terceira parte explora outro problema real do BarberSaaS: a necessidade de fortalecer a autenticacao e garantir que cada perfil acesse somente aquilo que lhe compete. O recorte considera o login, a sessao por cookie, a protecao de paginas e APIs e a associacao correta entre usuario autenticado e perfil operacional.",
    )

    base.add_heading(document, "1. PDCAs aplicados ao problema de autenticacao", 1)

    base.add_heading(document, "1.1 PDCA - Login e validade da sessao", 2)
    base.add_bullets(
        document,
        [
            "Plan: definir regras claras para login, cookie seguro, expiracao e sessao valida.",
            "Do: validar credenciais, emissao do token e redirecionamento por perfil.",
            "Check: observar erros de autenticacao, expiracao e sessao invalida.",
            "Act: ajustar politicas, mensagens e seguranca do fluxo de acesso.",
        ],
    )
    base.add_image(document, image_paths["pdca_1"], "Figura 1 - PDCA para login e validade da sessao")

    base.add_heading(document, "1.2 PDCA - Controle de acesso por perfil", 2)
    base.add_bullets(
        document,
        [
            "Plan: garantir que CLIENT, BARBER e ADMIN tenham acessos separados e coerentes.",
            "Do: revisar proxy, regras de rota e autorizacao nas APIs protegidas.",
            "Check: conferir tentativas indevidas, retornos 401 e 403 e redirecionamentos.",
            "Act: corrigir inconsistencias de permissao e formalizar criterios por perfil.",
        ],
    )
    base.add_image(document, image_paths["pdca_2"], "Figura 2 - PDCA para controle de acesso por perfil")

    base.add_heading(document, "1.3 PDCA - Vinculacao entre usuario e perfil operacional", 2)
    base.add_bullets(
        document,
        [
            "Plan: evitar que barbeiros ou administradores acessem dados errados por mapeamento incorreto.",
            "Do: validar a resolucao do barberId e o comportamento das rotas internas.",
            "Check: analisar erros de associacao, falta de perfil e bloqueios indevidos.",
            "Act: reforcar consistencia cadastral e rastreabilidade do vinculo entre usuario e modulo.",
        ],
    )
    base.add_image(document, image_paths["pdca_3"], "Figura 3 - PDCA para vinculacao entre usuario e perfil")

    base.add_heading(document, "2. Diagramas de causa e efeito", 1)

    base.add_heading(document, "2.1 Acesso indevido ou bloqueio incorreto de perfil", 2)
    base.add_body_paragraph(
        document,
        "Este diagrama trata o problema mais visivel para a equipe: o usuario entra no sistema, mas acessa uma area errada ou recebe bloqueio mesmo possuindo permissao. Isso afeta tanto a seguranca quanto a usabilidade dos modulos internos.",
    )
    base.add_image(document, image_paths["fishbone_1"], "Figura 4 - Causa e efeito para acesso indevido ou bloqueio incorreto")

    base.add_heading(document, "2.2 Falha de login ou sessao invalida", 2)
    base.add_body_paragraph(
        document,
        "Neste caso, o foco esta no momento de entrada do usuario. Credencial incorreta, cookie mal configurado, expiracao precoce ou assinatura invalida tornam o acesso instavel e reduzem a confianca no sistema.",
    )
    base.add_image(document, image_paths["fishbone_2"], "Figura 5 - Causa e efeito para falha de login ou sessao invalida")

    base.add_heading(document, "2.3 Barbeiro autenticado sem acesso ao proprio contexto", 2)
    base.add_body_paragraph(
        document,
        "Este terceiro diagrama aborda um problema mais especifico do dominio: o usuario pode estar autenticado, mas sem associacao correta ao barbeiro ativo. Isso compromete agenda, bloqueios e operacao diaria do profissional.",
    )
    base.add_image(document, image_paths["fishbone_3"], "Figura 6 - Causa e efeito para falha de vinculacao do barbeiro autenticado")

    base.add_heading(document, "3. Importancia das ferramentas", 1)
    base.add_body_paragraph(
        document,
        "O diagrama de causa e efeito ajuda a equipe a entender que problemas de autenticacao raramente nascem de uma unica origem. Uma falha de acesso pode envolver sessao, cookie, regras de rota, dados de usuario, mapeamento de perfil e ate configuracao de ambiente. Organizar essas causas melhora a qualidade da analise e reduz correcoes superficiais.",
    )
    base.add_body_paragraph(
        document,
        "O PDCA e importante porque transforma a seguranca e a coerencia de acesso em um processo de melhoria continua. A equipe planeja criterios de autenticacao, executa os fluxos, confere evidencias reais de uso e age sobre os desvios encontrados, fortalecendo tanto a protecao do sistema quanto a experiencia dos usuarios internos.",
    )

    document.save(base.DOCX_PATH)


def main() -> None:
    base.ensure_dirs()

    image_paths = {
        "pdca_1": base.create_pdca_image(
            "pdca_login_sessao.png",
            "PDCA 1 - Login e validade da sessao",
            "Foco",
            "Acesso confiavel e sessao valida",
            [
                "Definir regras para credenciais, cookie e expiracao.",
                "Garantir sessao assinada e redirecionamento correto por perfil.",
            ],
            [
                "Executar validacao de email, senha e status da conta.",
                "Emitir token de sessao e cookie httpOnly com tempo controlado.",
            ],
            [
                "Observar erros de login, sessao invalida e expiracao.",
                "Conferir consistencia do fluxo entre entrar e navegar nas areas protegidas.",
            ],
            [
                "Ajustar mensagens, seguranca e politicas de sessao.",
                "Padronizar verificacoes para evitar regressao no acesso.",
            ],
        ),
        "pdca_2": base.create_pdca_image(
            "pdca_controle_acesso_perfil.png",
            "PDCA 2 - Controle de acesso por perfil",
            "Foco",
            "CLIENT, BARBER e ADMIN com permissoes coerentes",
            [
                "Definir quais rotas e APIs pertencem a cada perfil.",
                "Estabelecer respostas esperadas para 401, 403 e redirecionamentos.",
            ],
            [
                "Revisar proxy, PAGE_RULES, API_RULES e autorizacao nas rotas.",
                "Executar testes de acesso indevido e navegacao autenticada.",
            ],
            [
                "Conferir se perfis errados sao barrados corretamente.",
                "Analisar redirecionamentos e retornos de sem permissao.",
            ],
            [
                "Corrigir inconsistencias e formalizar criterios por papel.",
                "Fortalecer a manutencao das regras de acesso.",
            ],
        ),
        "pdca_3": base.create_pdca_image(
            "pdca_vinculacao_perfil_operacional.png",
            "PDCA 3 - Vinculacao entre usuario e perfil operacional",
            "Foco",
            "Associacao correta entre conta e contexto do barbeiro",
            [
                "Definir regra segura para resolver o barberId do usuario.",
                "Evitar acesso sem perfil operacional valido.",
            ],
            [
                "Validar mapeamento por userId e fallback por email.",
                "Executar fluxos de agenda, bloqueios e disponibilidade do barbeiro.",
            ],
            [
                "Observar erros de perfil nao encontrado e bloqueios indevidos.",
                "Comparar contexto da sessao com o barbeiro realmente vinculado.",
            ],
            [
                "Corrigir inconsistencias cadastrais e melhorar rastreabilidade.",
                "Reduzir risco de acesso ao contexto errado.",
            ],
        ),
        "fishbone_1": base.create_fishbone_image(
            "causa_efeito_acesso_indevido.png",
            "Diagrama 1 - Acesso indevido ou bloqueio incorreto de perfil",
            "Acesso indevido ou bloqueio",
            [
                ("Regras", ["pathPrefix mal definido", "perfil permitido incorretamente"]),
                ("Sessao", ["role divergente do usuario", "token antigo ou inconsistente"]),
                ("Interface", ["rota errada apos login", "navegacao sem contexto correto"]),
            ],
            [
                ("API", ["autorizacao incompleta", "retorno 401 e 403 mal tratado"]),
                ("Dados", ["usuario inativo", "papel cadastrado incorretamente"]),
                ("Processo", ["criterios por perfil pouco claros", "teste insuficiente de permissao"]),
            ],
        ),
        "fishbone_2": base.create_fishbone_image(
            "causa_efeito_login_sessao.png",
            "Diagrama 2 - Falha de login ou sessao invalida",
            "Falha de login ou sessao",
            [
                ("Credencial", ["senha incorreta", "email digitado errado"]),
                ("Cookie", ["secure mal configurado", "cookie ausente ou expirado"]),
                ("Token", ["assinatura invalida", "expiracao antes do esperado"]),
            ],
            [
                ("Conta", ["usuario inativo", "papel divergente da sessao"]),
                ("Ambiente", ["SESSION_SECRET inadequado", "APP_BASE_URL inconsistente"]),
                ("Experiencia", ["mensagem pouco clara", "usuario nao entende o motivo da falha"]),
            ],
        ),
        "fishbone_3": base.create_fishbone_image(
            "causa_efeito_vinculacao_barbeiro.png",
            "Diagrama 3 - Barbeiro autenticado sem acesso ao proprio contexto",
            "Falha na vinculacao do barbeiro",
            [
                ("Cadastro", ["userId nao associado", "email do barbeiro divergente"]),
                ("Resolucao", ["fallback por email insuficiente", "perfil ativo nao localizado"]),
                ("Operacao", ["agenda e bloqueios sem contexto", "acesso parcial ao modulo"]),
            ],
            [
                ("Dados", ["status inativo do barbeiro", "duplicidade ou inconsistencia cadastral"]),
                ("Permissao", ["sessao valida mas contexto invalido", "barreiras indevidas na API"]),
                ("Qualidade", ["falta de teste de vinculo", "erro descoberto so em uso real"]),
            ],
        ),
    }

    build_docx_part3(image_paths)
    print(f"Imagens geradas em: {base.OUT_DIR}")
    print(f"Arquivo Word gerado em: {base.DOCX_PATH}")


if __name__ == "__main__":
    main()
