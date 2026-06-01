from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs_assets" / "quality_exercise"
DOCX_PATH = ROOT / "LISTA-GP-QUALIDADE-COM-IMAGENS.docx"

CANVAS_W = 1800
CANVAS_H = 1100

BG = "#f7f4ef"
CARD = "#fffdfa"
TEXT = "#201d18"
MUTED = "#6b655e"
PRIMARY = "#c59f59"
PRIMARY_DARK = "#7c5b24"
BORDER = "#ddd3c4"
ACCENT_1 = "#f4ead8"
ACCENT_2 = "#efe4d2"
ACCENT_3 = "#ede7dc"
ACCENT_4 = "#f3efe7"


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False):
    candidates = []
    if bold:
        candidates.extend(
            [
                Path(r"C:\Windows\Fonts\arialbd.ttf"),
                Path(r"C:\Windows\Fonts\calibrib.ttf"),
            ]
        )
    else:
        candidates.extend(
            [
                Path(r"C:\Windows\Fonts\arial.ttf"),
                Path(r"C:\Windows\Fonts\calibri.ttf"),
            ]
        )

    for path in candidates:
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue

    return ImageFont.load_default()


def text_wrap(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    if not words:
        return []

    lines: list[str] = []
    current = words[0]

    for word in words[1:]:
        trial = f"{current} {word}"
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word

    lines.append(current)
    return lines


def draw_multiline(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    lines: Iterable[str],
    font,
    fill: str,
    line_gap: int,
) -> None:
    x, y = xy
    for line in lines:
        draw.text((x, y), line, fill=fill, font=font)
        bbox = draw.textbbox((x, y), line, font=font)
        y += (bbox[3] - bbox[1]) + line_gap


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: list[str],
    fill: str,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=28, fill=fill, outline=BORDER, width=3)
    draw.text((x1 + 26, y1 + 22), title, fill=PRIMARY_DARK, font=get_font(28, bold=True))

    body_font = get_font(20)
    current_y = y1 + 74
    for item in body:
        wrapped = text_wrap(draw, item, body_font, x2 - x1 - 64)
        bullet_lines = []
        for idx, line in enumerate(wrapped):
            prefix = "• " if idx == 0 else "  "
            bullet_lines.append(f"{prefix}{line}")
        draw_multiline(draw, (x1 + 28, current_y), bullet_lines, body_font, TEXT, 6)
        line_height = draw.textbbox((0, 0), "Ag", font=body_font)[3]
        current_y += (line_height + 6) * len(bullet_lines) + 12


def create_canvas(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (CANVAS_W, CANVAS_H), BG)
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((36, 36, CANVAS_W - 36, CANVAS_H - 36), radius=34, fill=BG, outline="#e8dccb", width=3)
    draw.text((92, 72), title, fill=TEXT, font=get_font(40, bold=True))
    draw.text((92, 128), subtitle, fill=MUTED, font=get_font(24))
    draw.line((92, 168, CANVAS_W - 92, 168), fill="#eadfce", width=3)

    return image, draw


def arrow_right(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = PRIMARY) -> None:
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=8)
    draw.polygon([(x2, y2), (x2 - 24, y2 - 16), (x2 - 24, y2 + 16)], fill=color)


def arrow_down(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = PRIMARY) -> None:
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=8)
    draw.polygon([(x2, y2), (x2 - 16, y2 - 24), (x2 + 16, y2 - 24)], fill=color)


def arrow_left(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = PRIMARY) -> None:
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=8)
    draw.polygon([(x2, y2), (x2 + 24, y2 - 16), (x2 + 24, y2 + 16)], fill=color)


def arrow_up(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = PRIMARY) -> None:
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=color, width=8)
    draw.polygon([(x2, y2), (x2 - 16, y2 + 24), (x2 + 16, y2 + 24)], fill=color)


def create_pdca_image(
    filename: str,
    title: str,
    center_title: str,
    center_subtitle: str,
    plan_items: list[str],
    do_items: list[str],
    check_items: list[str],
    act_items: list[str],
) -> Path:
    image, draw = create_canvas(title, "Ciclo PDCA aplicado ao projeto BarberSaaS")

    top_box = (630, 220, 1170, 430)
    right_box = (1220, 430, 1730, 670)
    bottom_box = (630, 730, 1170, 940)
    left_box = (70, 430, 580, 670)

    draw_box(draw, top_box, "PLAN", plan_items, ACCENT_1)
    draw_box(draw, right_box, "DO", do_items, ACCENT_2)
    draw_box(draw, bottom_box, "CHECK", check_items, ACCENT_3)
    draw_box(draw, left_box, "ACT", act_items, ACCENT_4)

    draw.ellipse((720, 450, 1080, 810), fill=CARD, outline=PRIMARY, width=8)
    draw.text((805, 535), center_title, fill=TEXT, font=get_font(34, bold=True))

    center_font = get_font(22)
    center_lines = text_wrap(draw, center_subtitle, center_font, 250)
    draw_multiline(draw, (793, 590), center_lines, center_font, MUTED, 6)

    arrow_right(draw, (1170, 390), (1320, 390))
    arrow_down(draw, (1480, 670), (1480, 760))
    arrow_left(draw, (630, 885), (480, 885))
    arrow_up(draw, (320, 430), (320, 340))

    out_path = OUT_DIR / filename
    image.save(out_path)
    return out_path


def create_fishbone_image(
    filename: str,
    title: str,
    effect: str,
    branches_top: list[tuple[str, list[str]]],
    branches_bottom: list[tuple[str, list[str]]],
) -> Path:
    image, draw = create_canvas(title, "Diagrama de causa e efeito sobre riscos e atividades do sistema")

    spine_y = 600
    draw.line((190, spine_y, 1450, spine_y), fill=PRIMARY, width=10)
    draw.polygon([(1450, 550), (1600, spine_y), (1450, 650)], fill=PRIMARY)
    draw.rounded_rectangle((1360, 395, 1730, 525), radius=26, fill=ACCENT_1, outline=PRIMARY, width=4)
    draw.text((1390, 430), effect, fill=TEXT, font=get_font(30, bold=True))

    upper_positions = [420, 760, 1100]
    lower_positions = [420, 760, 1100]
    label_font = get_font(24, bold=True)
    item_font = get_font(18)

    for idx, (label, items) in enumerate(branches_top):
        x = upper_positions[idx]
        draw.line((x, spine_y, x - 140, 320), fill="#8c7a60", width=5)
        draw.text((x - 295, 268), label, fill=PRIMARY_DARK, font=label_font)
        for item_idx, item in enumerate(items):
            y = 316 + item_idx * 44
            draw.line((x - 52, 472 - item_idx * 22, x - 100, y + 12), fill="#8c7a60", width=3)
            wrapped = text_wrap(draw, item, item_font, 220)
            for wrap_idx, wrap_line in enumerate(wrapped[:2]):
                prefix = "• " if wrap_idx == 0 else "  "
                draw.text((x - 310, y + wrap_idx * 22), f"{prefix}{wrap_line}", fill=TEXT, font=item_font)

    for idx, (label, items) in enumerate(branches_bottom):
        x = lower_positions[idx]
        draw.line((x, spine_y, x - 140, 880), fill="#8c7a60", width=5)
        draw.text((x - 295, 890), label, fill=PRIMARY_DARK, font=label_font)
        for item_idx, item in enumerate(items):
            y = 932 + item_idx * 44
            draw.line((x - 52, 726 + item_idx * 22, x - 100, y + 10), fill="#8c7a60", width=3)
            wrapped = text_wrap(draw, item, item_font, 220)
            for wrap_idx, wrap_line in enumerate(wrapped[:2]):
                prefix = "• " if wrap_idx == 0 else "  "
                draw.text((x - 310, y + wrap_idx * 22), f"{prefix}{wrap_line}", fill=TEXT, font=item_font)

    out_path = OUT_DIR / filename
    image.save(out_path)
    return out_path


def set_run_font(run, size: int, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    set_run_font(run, max(14, 20 - level * 2), bold=True)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, 12)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_run_font(run, 12)


def add_image(document: Document, path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(16.5))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, 11, bold=True, color=PRIMARY_DARK)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def build_docx(image_paths: dict[str, Path]) -> None:
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Lista de Exercicios - GP Qualidade")
    set_run_font(run, 18, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("BarberSaaS - PDCAs e Diagramas de Causa e Efeito com imagens")
    set_run_font(run, 12, color=PRIMARY_DARK)

    add_body_paragraph(
        document,
        "Este material foi elaborado com base no sistema BarberSaaS, considerando o fluxo de agendamento, a validacao de disponibilidade, as notificacoes externas, o controle de acesso por perfil e os riscos descritos na documentacao tecnica do projeto.",
    )

    add_heading(document, "1. PDCAs do projeto", 1)

    add_heading(document, "1.1 PDCA - Qualidade do fluxo de agendamento", 2)
    add_bullets(
        document,
        [
            "Plan: definir criterios para mostrar apenas horarios realmente livres ao cliente.",
            "Do: executar validacoes na logica de disponibilidade e na criacao do agendamento.",
            "Check: conferir conflitos, erros de API e cenarios de borda.",
            "Act: ajustar regras, ampliar testes e formalizar criterios de aceite.",
        ],
    )
    add_image(document, image_paths["pdca_1"], "Figura 1 - PDCA do fluxo de agendamento")

    add_heading(document, "1.2 PDCA - Estabilidade do sistema e configuracao", 2)
    add_bullets(
        document,
        [
            "Plan: mapear riscos ligados a banco, variaveis de ambiente e integracoes externas.",
            "Do: testar o ambiente e validar os fallbacks de notificacao.",
            "Check: analisar logs, respostas de API e continuidade do agendamento sem dependencias externas.",
            "Act: padronizar configuracoes, reforcar monitoramento e criar checklists de entrega.",
        ],
    )
    add_image(document, image_paths["pdca_2"], "Figura 2 - PDCA da estabilidade tecnica do sistema")

    add_heading(document, "1.3 PDCA - Integracao dos modulos internos e acesso por perfil", 2)
    add_bullets(
        document,
        [
            "Plan: garantir que cliente, barbeiro e administrador usem dados e permissoes corretos.",
            "Do: revisar rotas protegidas, redirecionamentos e consumo de dados nas telas internas.",
            "Check: validar se cada perfil acessa apenas os recursos apropriados.",
            "Act: priorizar correcoes de permissao, integracao e criterios de aceite por perfil.",
        ],
    )
    add_image(document, image_paths["pdca_3"], "Figura 3 - PDCA da consolidacao dos modulos internos")

    add_heading(document, "2. Diagramas de causa e efeito", 1)

    add_heading(document, "2.1 Espinha de peixe - Falha no agendamento", 2)
    add_body_paragraph(
        document,
        "O efeito central deste diagrama e a perda de confiabilidade no principal fluxo do produto. Como o agendamento e o nucleo do BarberSaaS, qualquer falha em regra, dado ou validacao impacta diretamente a experiencia do cliente e a organizacao da barbearia.",
    )
    add_image(document, image_paths["fishbone_1"], "Figura 4 - Causa e efeito para conflito ou falha no agendamento")

    add_heading(document, "2.2 Espinha de peixe - Indisponibilidade do sistema", 2)
    add_body_paragraph(
        document,
        "Neste caso, a indisponibilidade e tratada como um risco tecnico e de processo. O diagrama mostra que uma falha pode nascer nao apenas da infraestrutura, mas tambem de configuracao incompleta, dependencia externa e monitoramento insuficiente.",
    )
    add_image(document, image_paths["fishbone_2"], "Figura 5 - Causa e efeito para indisponibilidade do sistema")

    add_heading(document, "2.3 Espinha de peixe - Atraso e queda da qualidade", 2)
    add_body_paragraph(
        document,
        "Este diagrama associa atraso de entrega e queda de qualidade a fatores como crescimento indevido do escopo, integracao incompleta, comunicacao falha, revisao tardia e priorizacao difusa entre os modulos do sistema.",
    )
    add_image(document, image_paths["fishbone_3"], "Figura 6 - Causa e efeito para atraso do projeto e perda de qualidade")

    add_heading(document, "3. Importancia das ferramentas", 1)
    add_body_paragraph(
        document,
        "O diagrama de causa e efeito e importante porque ajuda a equipe a sair da analise superficial do problema e organizar as possiveis origens de cada falha em grupos logicos, como metodo, pessoas, dados, infraestrutura e processo. Isso melhora a tomada de decisao e torna mais claro o que precisa ser tratado de forma preventiva.",
    )
    add_body_paragraph(
        document,
        "O PDCA e importante porque conecta planejamento, execucao, verificacao e acao corretiva em um fluxo continuo de melhoria. No BarberSaaS, essa logica ajuda a proteger o nucleo de agendamento, fortalecer a estabilidade do ambiente e amadurecer os modulos internos com base em evidencias observadas durante o desenvolvimento.",
    )

    document.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()

    image_paths = {
        "pdca_1": create_pdca_image(
            "pdca_fluxo_agendamento.png",
            "PDCA 1 - Qualidade do fluxo de agendamento",
            "Foco",
            "Confianca no agendamento online",
            [
                "Definir meta de mostrar apenas horarios realmente livres.",
                "Considerar disponibilidade, bloqueios, duracao e agendamentos ativos.",
            ],
            [
                "Executar validacoes em availability.ts e na criacao do agendamento.",
                "Testar servicos com duracoes diferentes e reagendamentos.",
            ],
            [
                "Conferir conflitos de horario e respostas da API.",
                "Revisar cenarios de borda e tentativas rejeitadas.",
            ],
            [
                "Ajustar regras com falha.",
                "Ampliar testes e documentar criterios de aceite.",
            ],
        ),
        "pdca_2": create_pdca_image(
            "pdca_estabilidade_sistema.png",
            "PDCA 2 - Estabilidade do sistema e configuracao",
            "Foco",
            "Reducao de falhas tecnicas e de ambiente",
            [
                "Mapear riscos em banco, sessao, e-mail e WhatsApp.",
                "Definir indicadores de erro, disponibilidade e continuidade do fluxo.",
            ],
            [
                "Executar checklist tecnico do ambiente.",
                "Testar o agendamento com e sem integracoes externas configuradas.",
            ],
            [
                "Analisar logs, respostas da API e comportamento de fallback.",
                "Confirmar que a reserva continua mesmo sem notificacao externa.",
            ],
            [
                "Padronizar configuracoes e reforcar monitoramento.",
                "Transformar a validacao do ambiente em rotina de entrega.",
            ],
        ),
        "pdca_3": create_pdca_image(
            "pdca_modulos_internos.png",
            "PDCA 3 - Modulos internos e seguranca por perfil",
            "Foco",
            "Integracao correta entre dados, telas e permissoes",
            [
                "Definir metas para cliente, barbeiro e administrador.",
                "Usar indicadores de permissao, integracao e acesso indevido.",
            ],
            [
                "Revisar rotas protegidas, sessoes e redirecionamentos.",
                "Conferir se cada tela consome apenas os dados apropriados.",
            ],
            [
                "Verificar se cada perfil acessa apenas o que lhe compete.",
                "Comparar paineis com os dados reais retornados pelas APIs.",
            ],
            [
                "Priorizar correcoes de integracao e permissao.",
                "Formalizar criterios de aceite por perfil de usuario.",
            ],
        ),
        "fishbone_1": create_fishbone_image(
            "causa_efeito_agendamento.png",
            "Diagrama 1 - Conflito ou falha no agendamento",
            "Falha no agendamento",
            [
                ("Metodo", ["criterios de aceite insuficientes", "teste manual excessivo"]),
                ("Sistema", ["erro no calculo de duracao", "conflito entre bloqueios e reservas"]),
                ("Dados", ["cadastro incorreto de servico", "disponibilidade inconsistente"]),
            ],
            [
                ("Pessoas", ["erro no cadastro operacional", "revisao funcional tardia"]),
                ("Infraestrutura", ["lentidao de banco ou API", "sessao expirada no fluxo"]),
                ("Medicao", ["falta de metricas sobre conflitos", "sem controle de tentativas rejeitadas"]),
            ],
        ),
        "fishbone_2": create_fishbone_image(
            "causa_efeito_indisponibilidade.png",
            "Diagrama 2 - Indisponibilidade total ou parcial do sistema",
            "Indisponibilidade",
            [
                ("Infraestrutura", ["banco fora do ar", "servidor ou rede instavel"]),
                ("Configuracao", ["variaveis ausentes", "credenciais invalidas"]),
                ("Dependencias", ["SMTP indisponivel", "API do WhatsApp falhando"]),
            ],
            [
                ("Seguranca", ["problemas de sessao", "regra de acesso incorreta"]),
                ("Processo", ["deploy sem checklist", "sem plano de contingencia"]),
                ("Monitoramento", ["logs insuficientes", "demora para achar causa raiz"]),
            ],
        ),
        "fishbone_3": create_fishbone_image(
            "causa_efeito_atraso_qualidade.png",
            "Diagrama 3 - Atraso na entrega e queda da qualidade",
            "Atraso e queda de qualidade",
            [
                ("Escopo", ["muitos modulos ao mesmo tempo", "requisitos crescendo sem recorte"]),
                ("Integracao", ["interface sem API consolidada", "ajustes frequentes entre camadas"]),
                ("Comunicacao", ["validacao tardia com o responsavel", "prioridades sem registro claro"]),
            ],
            [
                ("Pessoas", ["distribuicao desigual de tarefas", "conhecimento tecnico concentrado"]),
                ("Validacao", ["revisao continua insuficiente", "correcao apenas no fim"]),
                ("Priorizacao", ["foco difuso entre modulos", "nucleo funcional pouco protegido"]),
            ],
        ),
    }

    build_docx(image_paths)
    print(f"Imagens geradas em: {OUT_DIR}")
    print(f"Arquivo Word gerado em: {DOCX_PATH}")


if __name__ == "__main__":
    main()
